from docx import Document

# 건위천 괘 샘플 파일 생성
doc = Document()

doc.add_heading('1. 건위천(乾爲天)', 0)

doc.add_paragraph('건위천은 64괘 중 첫 번째 괘로, 순수한 양의 에너지를 나타냅니다.')

doc.add_heading('괘사(卦辭)', level=1)
doc.add_paragraph('건(乾): 원형이정(元亨利貞)')
doc.add_paragraph('건괘는 크게 형통하니 올바름이 이롭다.')

doc.add_heading('상전(象傳)', level=1)
doc.add_paragraph('천행건 군자이자강불식(天行健 君子以自彊不息)')
doc.add_paragraph('하늘의 운행이 건실하니, 군자는 스스로 힘써 쉬지 않는다.')

doc.add_heading('해석', level=1)
doc.add_paragraph('건위천 괘는 창조와 리더십을 상징합니다. 모든 일이 잘 풀릴 수 있는 길한 괘이지만, '
                 '지속적인 노력과 올바른 방향성이 필요합니다.')

doc.add_paragraph('이 괘가 나왔다면:')
doc.add_paragraph('• 새로운 시작에 좋은 때입니다')
doc.add_paragraph('• 리더십을 발휘할 기회가 옵니다')
doc.add_paragraph('• 꾸준한 노력이 성과로 이어집니다')
doc.add_paragraph('• 정직하고 올바른 길을 걸으세요')

doc.save('해석/건위천.docx')
print("건위천.docx 파일이 생성되었습니다.")

# 곤위지 괘 샘플 파일도 생성
doc2 = Document()

doc2.add_heading('2. 곤위지(坤爲地)', 0)

doc2.add_paragraph('곤위지는 64괘 중 두 번째 괘로, 순수한 음의 에너지를 나타냅니다.')

doc2.add_heading('괘사(卦辭)', level=1)
doc2.add_paragraph('곤(坤): 원형 우말이정(元亨 牝馬之貞)')
doc2.add_paragraph('곤괘는 원대하게 형통하니 암말의 정절과 같아야 한다.')

doc2.add_heading('상전(象傳)', level=1)
doc2.add_paragraph('지세곤 군자이후덕재물(地勢坤 君子以厚德載物)')
doc2.add_paragraph('땅의 형세가 곤순하니, 군자는 두터운 덕으로 만물을 포용한다.')

doc2.add_heading('해석', level=1)
doc2.add_paragraph('곤위지 괘는 포용과 순응을 상징합니다. 겸손하고 온유한 자세로 상황을 받아들이며, '
                 '남을 받쳐주는 역할을 할 때 좋은 결과를 얻을 수 있습니다.')

doc2.add_paragraph('이 괘가 나왔다면:')
doc2.add_paragraph('• 겸손한 자세가 필요합니다')
doc2.add_paragraph('• 다른 사람을 받쳐주는 역할을 하세요')
doc2.add_paragraph('• 순응하며 때를 기다리세요')
doc2.add_paragraph('• 포용력을 발휘할 때입니다')

doc2.save('해석/곤위지.docx')
print("곤위지.docx 파일이 생성되었습니다.") 