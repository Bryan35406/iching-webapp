import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
import os
import subprocess
import platform
from pathlib import Path
import docx
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from datetime import datetime
import random
import re

class IChingReader:
    def __init__(self, root):
        self.root = root
        self.root.title("주역 괘 해석 프로그램")
        self.root.geometry("1300x950")
        self.root.configure(bg='#f0f0f0')
        
        # 글자 크기 설정 변수 (기본값을 아주 크게로 설정)
        self.font_sizes = {
            'title': 27,      # 18 * 1.5
            'subtitle': 18,   # 12 * 1.5
            'label': 15,      # 10 * 1.5
            'text': 14,       # 9 * 1.5 (약간 올림)
            'entry': 18       # 12 * 1.5
        }
        
        # 현재 점괘 결과 저장용 변수들
        self.current_question = ""
        self.current_lines = []
        self.current_moving_lines = []
        self.current_original = None
        self.current_changed = None
        self.current_final = None
        self.current_rule = ""
        self.current_highlight_lines = []  # 강조할 효 번호들
        
        # 64괘 정보 (괘번호: (괘이름, 이진코드)) - 정확한 전통 패턴
        self.hexagrams = {
            1: ("1괘", "111111"),
            2: ("2괘", "222222"),
            3: ("3괘", "122212"),
            4: ("4괘", "212221"),
            5: ("5괘", "111212"),
            6: ("6괘", "212111"),
            7: ("7괘", "212222"),
            8: ("8괘", "222212"),
            9: ("9괘", "111211"),
            10: ("10괘", "112111"),
            11: ("11괘", "111222"),
            12: ("12괘", "222111"),
            13: ("13괘", "121111"),
            14: ("14괘", "111121"),
            15: ("15괘", "221222"),
            16: ("16괘", "222122"),
            17: ("17괘", "122112"),
            18: ("18괘", "211221"),
            19: ("19괘", "112222"),
            20: ("20괘", "222211"),
            21: ("21괘", "122121"),
            22: ("22괘", "121221"),
            23: ("23괘", "222221"),
            24: ("24괘", "122222"),
            25: ("25괘", "122111"),
            26: ("26괘", "111221"),
            27: ("27괘", "122221"),
            28: ("28괘", "211112"),
            29: ("29괘", "212212"),
            30: ("30괘", "121121"),
            31: ("31괘", "221112"),
            32: ("32괘", "211122"),
            33: ("33괘", "221111"),
            34: ("34괘", "111122"),
            35: ("35괘", "222121"),
            36: ("36괘", "121222"),
            37: ("37괘", "121211"),
            38: ("38괘", "112121"),
            39: ("39괘", "221212"),
            40: ("40괘", "212122"),
            41: ("41괘", "112221"),
            42: ("42괘", "122211"),
            43: ("43괘", "111112"),
            44: ("44괘", "211111"),
            45: ("45괘", "222112"),
            46: ("46괘", "211222"),
            47: ("47괘", "212112"),
            48: ("48괘", "211212"),
            49: ("49괘", "121112"),
            50: ("50괘", "211121"),
            51: ("51괘", "122122"),
            52: ("52괘", "221221"),
            53: ("53괘", "221211"),
            54: ("54괘", "112122"),
            55: ("55괘", "121122"),
            56: ("56괘", "221121"),
            57: ("57괘", "211211"),
            58: ("58괘", "112112"),
            59: ("59괘", "212211"),
            60: ("60괘", "112212"),
            61: ("61괘", "112211"),
            62: ("62괘", "221122"),
            63: ("63괘", "121212"),
            64: ("64괘", "212121")
        }
        
        self.setup_menu()
        self.setup_ui()
    
    def setup_menu(self):
        """메뉴바 설정"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # 파일 메뉴
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="파일", menu=file_menu)
        file_menu.add_command(label="점괘 결과 저장", command=self.save_result)
        file_menu.add_separator()
        file_menu.add_command(label="종료", command=self.root.quit)
        
        # 도구 메뉴
        tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="도구", menu=tools_menu)
        tools_menu.add_command(label="자동 괘 뽑기", command=self.auto_generate_hexagram)
        
        # 설정 메뉴
        settings_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="설정", menu=settings_menu)
        
        # 글자 크기 서브메뉴
        font_menu = tk.Menu(settings_menu, tearoff=0)
        settings_menu.add_cascade(label="글자 크기", menu=font_menu)
        
        font_menu.add_command(label="아주 작게", command=lambda: self.change_font_size(0.7))
        font_menu.add_command(label="작게", command=lambda: self.change_font_size(0.85))
        font_menu.add_command(label="보통", command=lambda: self.change_font_size(1.0))
        font_menu.add_command(label="크게", command=lambda: self.change_font_size(1.2))
        font_menu.add_command(label="아주 크게 ✓", command=lambda: self.change_font_size(1.5))
        
        # 도움말 메뉴
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="도움말", menu=help_menu)
        help_menu.add_command(label="사용법", command=self.show_help)
        help_menu.add_command(label="프로그램 정보", command=self.show_about)
    
    def auto_generate_hexagram(self):
        """자동으로 괘와 동효를 생성"""
        response = messagebox.askyesno("자동 괘 뽑기", 
                                     "자동으로 괘를 뽑으시겠습니까?\n"
                                     "현재 입력된 내용이 모두 지워집니다.")
        if not response:
            return
        
        # 기존 입력 지우기
        self.clear_inputs()
        
        # 6개 효 랜덤 생성 (1=양, 2=음)
        lines = []
        for i in range(6):
            line_value = random.choice(['1', '2'])
            lines.append(line_value)
            self.line_vars[i].set(line_value)
        
        # 동효 랜덤 생성 (실제 동전 던지기 확률 적용)
        # 실제 확률: 0개(17.8%), 1개(35.6%), 2개(29.7%), 3개(13.2%), 4개(3.3%), 5개(0.4%), 6개(0.02%)
        moving_count = random.choices(
            population=[0, 1, 2, 3, 4, 5, 6], 
            weights=[17.8, 35.6, 29.7, 13.2, 3.3, 0.4, 0.02], 
            k=1
        )[0]
        
        if moving_count > 0:
            moving_positions = random.sample(range(6), moving_count)
            for pos in moving_positions:
                self.moving_vars[pos].set(True)
        
        # 자동 식별
        self.identify_hexagram()
        
        messagebox.showinfo("자동 괘 뽑기 완료", 
                          f"자동으로 괘가 생성되었습니다!\n"
                          f"효: {' '.join(lines)}\n"
                          f"동효 개수: {moving_count}개\n"
                          f"(실제 동전 던지기 확률 적용)")
    
    def change_font_size(self, scale):
        """글자 크기 변경"""
        base_sizes = {
            'title': 18,
            'subtitle': 12,
            'label': 10,
            'text': 9,
            'entry': 12
        }
        
        # 새로운 크기 계산
        for key in self.font_sizes:
            self.font_sizes[key] = max(8, int(base_sizes[key] * scale))
        
        # UI 업데이트
        self.update_fonts()
    
    def update_fonts(self):
        """모든 위젯의 폰트 업데이트"""
        try:
            # 제목
            self.title_label.config(font=('맑은 고딕', self.font_sizes['title'], 'bold'))
            
            # 설명
            self.desc_label.config(font=('맑은 고딕', self.font_sizes['subtitle']))
            
            # 질문 입력 필드
            self.question_entry.config(font=('맑은 고딕', self.font_sizes['entry']))
            
            # 입력 필드들
            for entry in self.line_vars:
                entry.config(font=('맑은 고딕', self.font_sizes['entry']))
            
            # 괘 이름 라벨들
            self.original_name_label.config(font=('맑은 고딕', self.font_sizes['subtitle'], 'bold'))
            self.changed_name_label.config(font=('맑은 고딕', self.font_sizes['subtitle'], 'bold'))
            self.final_name_label.config(font=('맑은 고딕', self.font_sizes['subtitle'], 'bold'))
            
            # 규칙 라벨
            self.rule_label.config(font=('맑은 고딕', self.font_sizes['label']))
            
            # 텍스트 위젯들
            self.original_text.config(font=('맑은 고딕', self.font_sizes['text']))
            self.changed_text.config(font=('맑은 고딕', self.font_sizes['text']))
            self.final_text.config(font=('맑은 고딕', self.font_sizes['text']))
            
            # 하이라이트 태그도 업데이트
            self.original_text.tag_configure("highlight", 
                background="#FFFF00", foreground="#FF0000", 
                font=('맑은 고딕', self.font_sizes['text'], 'bold'))
            self.changed_text.tag_configure("highlight", 
                background="#FFFF00", foreground="#FF0000", 
                font=('맑은 고딕', self.font_sizes['text'], 'bold'))
            self.final_text.tag_configure("highlight", 
                background="#FFFF00", foreground="#FF0000", 
                font=('맑은 고딕', self.font_sizes['text'], 'bold'))
                
        except AttributeError:
            # 아직 UI가 생성되지 않은 경우 무시
            pass
    
    def show_help(self):
        """사용법 도움말"""
        help_text = """
주역 괘 해석 프로그램 사용법

1. 질문 입력:
   - 점을 치고자 하는 질문을 입력하세요
   - 구체적이고 명확한 질문일수록 좋습니다

2. 효 입력:
   - 1효부터 6효까지 1(양) 또는 2(음)을 입력하세요
   - 동전이나 시초 던지기 결과를 입력합니다

3. 동효 선택:
   - 변하는 효에 "동효" 체크박스를 체크하세요
   - 동효 개수에 따라 해석 방법이 달라집니다

4. 동효별 해석 규칙:
   - 0개: 본괘 전체 해석
   - 1개: 본괘 + 동효 강조
   - 2개: 본괘 + 큰 번호 효 강조
   - 3개: 하괘/상괘 중심으로 본괘 또는 지괘
   - 4개: 지괘 + 고정효 중 작은 것 강조
   - 5개: 지괘 + 고정효 강조
   - 6개: 지괘 전체 해석

5. 자동 괘 뽑기:
   - 실제 동전 던지기 확률을 적용합니다
   - 동효 확률: 0개(17.8%), 1개(35.6%), 2개(29.7%), 
     3개(13.2%), 4개(3.3%), 5개(0.4%), 6개(0.02%)
   - 동효 있을 확률: 82.2%

6. 점괘 결과 저장:
   - 파일 > 점괘 결과 저장으로 결과를 워드 파일로 저장
   - 출력 폴더에 날짜별로 저장됩니다

7. 글자 크기:
   - 설정 > 글자 크기 메뉴에서 조절 가능

8. 결과:
   - 본괘, 지괘, 최종해석이 각각 표시됩니다
   - 중요한 효는 노란색으로 강조됩니다
        """
        
        help_window = tk.Toplevel(self.root)
        help_window.title("사용법")
        help_window.geometry("500x700")
        
        text_widget = scrolledtext.ScrolledText(help_window, wrap=tk.WORD, 
                                              font=('맑은 고딕', 11))
        text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        text_widget.insert(tk.END, help_text)
        text_widget.config(state=tk.DISABLED)
    
    def show_about(self):
        """프로그램 정보"""
        about_text = """
주역 괘 해석 프로그램 v1.1

전통적인 주역 점괘 해석을 위한 GUI 프로그램입니다.

기능:
• 질문 입력 및 기록
• 64괘 자동 식별
• 동효(변효) 처리
• 본괘/지괘/최종해석 표시
• 워드 파일 기반 해석 시스템
• 점괘 결과 저장 및 출력
• 글자 크기 조절
• 효별 색상 강조

개발: AI Assistant
버전: 1.1
        """
        
        messagebox.showinfo("프로그램 정보", about_text)
    
    def setup_ui(self):
        # 메인 프레임
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 제목
        self.title_label = ttk.Label(main_frame, text="주역 괘 해석 프로그램", 
                               font=('맑은 고딕', self.font_sizes['title'], 'bold'))
        self.title_label.grid(row=0, column=0, columnspan=4, pady=(0, 30))
        
        # 질문 입력 프레임
        question_frame = ttk.LabelFrame(main_frame, text="질문 입력", padding="15")
        question_frame.grid(row=1, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=(0, 25))
        
        ttk.Label(question_frame, text="점을 치고자 하는 질문:").grid(row=0, column=0, padx=(0, 10), sticky=tk.W)
        self.question_entry = ttk.Entry(question_frame, width=80, 
                                       font=('맑은 고딕', self.font_sizes['entry']))
        self.question_entry.grid(row=0, column=1, sticky=(tk.W, tk.E))
        question_frame.columnconfigure(1, weight=1)
        
        # 설명
        self.desc_label = ttk.Label(main_frame, text="1효부터 6효까지 음양을 선택하고 동효(변효)를 체크하세요", 
                              font=('맑은 고딕', self.font_sizes['subtitle']))
        self.desc_label.grid(row=2, column=0, columnspan=4, pady=(0, 20))
        
        # 효 입력 프레임
        input_frame = ttk.LabelFrame(main_frame, text="효 입력 및 동효 선택", padding="15")
        input_frame.grid(row=3, column=0, columnspan=4, sticky=(tk.W, tk.E), pady=(0, 25))
        
        self.line_vars = []  # 효 선택 변수들 (1=양, 2=음)
        self.moving_vars = []  # 동효 체크박스 변수들
        
        for i in range(6):
            # 효 번호 라벨 (가운데 정렬)
            ttk.Label(input_frame, text=f"{i+1}효:").grid(row=0, column=i*4, padx=(10, 5))
            
            # 음양 선택 라디오 버튼
            line_var = tk.StringVar(value="1")  # 기본값: 양
            self.line_vars.append(line_var)
            
            # 양 버튼 (가운데 정렬)
            yang_btn = ttk.Radiobutton(input_frame, text="양(―)", variable=line_var, value="1")
            yang_btn.grid(row=1, column=i*4, padx=(10, 2))
            
            # 음 버튼 (가운데 정렬)
            yin_btn = ttk.Radiobutton(input_frame, text="음(- -)", variable=line_var, value="2")
            yin_btn.grid(row=2, column=i*4, padx=(10, 2))
            
            # 동효 체크박스 (가운데 정렬, 위쪽에 간격 추가)
            moving_var = tk.BooleanVar()
            self.moving_vars.append(moving_var)
            check = ttk.Checkbutton(input_frame, text="동효", variable=moving_var)
            check.grid(row=3, column=i*4, padx=(10, 20), pady=(10, 0))
        
        # 버튼 프레임
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=4, pady=(0, 30))
        
        identify_btn = ttk.Button(button_frame, text="괘 식별하기", 
                                 command=self.identify_hexagram, 
                                 style='Accent.TButton')
        identify_btn.pack(side=tk.LEFT, padx=(0, 15))
        
        auto_btn = ttk.Button(button_frame, text="자동 괘 뽑기", 
                             command=self.auto_generate_hexagram)
        auto_btn.pack(side=tk.LEFT, padx=(0, 15))
        
        clear_btn = ttk.Button(button_frame, text="지우기", command=self.clear_inputs)
        clear_btn.pack(side=tk.LEFT, padx=(0, 15))
        
        save_btn = ttk.Button(button_frame, text="점괘 결과 저장", 
                             command=self.save_result)
        save_btn.pack(side=tk.LEFT)
        
        # 결과를 3등분하는 프레임
        result_main_frame = ttk.Frame(main_frame)
        result_main_frame.grid(row=5, column=0, columnspan=4, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 본괘 프레임 (왼쪽)
        original_frame = ttk.LabelFrame(result_main_frame, text="본괘", padding="15")
        original_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 8))
        
        # 본괘 이름 표시
        self.original_name_label = ttk.Label(original_frame, text="", 
                                           font=('맑은 고딕', self.font_sizes['subtitle'], 'bold'),
                                           foreground='#2E8B57')
        self.original_name_label.pack(pady=(0, 15))
        
        # 본괘 해석 텍스트
        self.original_text = scrolledtext.ScrolledText(original_frame, 
                                                     width=25, height=18,
                                                     font=('맑은 고딕', self.font_sizes['text']),
                                                     wrap=tk.WORD)
        self.original_text.pack(fill=tk.BOTH, expand=True)
        
        # 지괘 프레임 (가운데)
        changed_frame = ttk.LabelFrame(result_main_frame, text="지괘 (동효 반전)", padding="15")
        changed_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(8, 8))
        
        # 지괘 이름 표시
        self.changed_name_label = ttk.Label(changed_frame, text="", 
                                          font=('맑은 고딕', self.font_sizes['subtitle'], 'bold'),
                                          foreground='#8B2E57')
        self.changed_name_label.pack(pady=(0, 15))
        
        # 지괘 해석 텍스트
        self.changed_text = scrolledtext.ScrolledText(changed_frame, 
                                                    width=25, height=18,
                                                    font=('맑은 고딕', self.font_sizes['text']),
                                                    wrap=tk.WORD)
        self.changed_text.pack(fill=tk.BOTH, expand=True)
        
        # 최종 해석 프레임 (오른쪽)
        final_frame = ttk.LabelFrame(result_main_frame, text="최종 해석", padding="15")
        final_frame.grid(row=0, column=2, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(8, 0))
        
        # 최종 괘 이름 표시
        self.final_name_label = ttk.Label(final_frame, text="", 
                                        font=('맑은 고딕', self.font_sizes['subtitle'], 'bold'),
                                        foreground='#8B4513')
        self.final_name_label.pack(pady=(0, 15))
        
        # 동효 규칙 설명
        self.rule_label = ttk.Label(final_frame, text="", 
                                   font=('맑은 고딕', self.font_sizes['label']),
                                   foreground='#666666')
        self.rule_label.pack(pady=(0, 15))
        
        # 최종 해석 텍스트
        self.final_text = scrolledtext.ScrolledText(final_frame, 
                                                  width=25, height=18,
                                                  font=('맑은 고딕', self.font_sizes['text']),
                                                  wrap=tk.WORD)
        self.final_text.pack(fill=tk.BOTH, expand=True)
        
        # 그리드 가중치 설정
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(5, weight=1)
        result_main_frame.columnconfigure(0, weight=1)
        result_main_frame.columnconfigure(1, weight=1)
        result_main_frame.columnconfigure(2, weight=1)
        result_main_frame.rowconfigure(0, weight=1)
    
    def identify_hexagram(self):
        # 질문 저장
        self.current_question = self.question_entry.get().strip()
        
        # 입력값 검증 및 수집
        lines = []
        for i, var in enumerate(self.line_vars):
            value = var.get()
            if not value:
                messagebox.showerror("입력 오류", f"{i+1}효를 입력해주세요.")
                return
            if value not in ['1', '2']:
                messagebox.showerror("입력 오류", f"{i+1}효는 1(양) 또는 2(음)만 입력 가능합니다.")
                return
            lines.append(value)
        
        # 현재 효 저장
        self.current_lines = lines.copy()
        
        # 본괘 찾기
        original_pattern = ''.join(lines)
        original_hexagram = self.find_hexagram(original_pattern)
        
        if original_hexagram:
            number, name = original_hexagram
            self.current_original = original_hexagram
            self.original_name_label.config(text=f"{number}번째 괘: {name}")
            self.load_interpretation(name, self.original_text)
        else:
            messagebox.showerror("오류", "해당하는 본괘를 찾을 수 없습니다.")
            return
        
        # 동효 확인 및 지괘 계산
        moving_lines = []
        for i, var in enumerate(self.moving_vars):
            if var.get():  # 체크된 동효
                moving_lines.append(i)
        
        # 현재 동효 저장
        self.current_moving_lines = moving_lines.copy()
        
        if moving_lines:
            # 동효가 있으면 지괘 계산
            changed_lines = lines.copy()
            for i in moving_lines:
                # 음양 반전 (1 -> 2, 2 -> 1)
                changed_lines[i] = '2' if changed_lines[i] == '1' else '1'
            
            changed_pattern = ''.join(changed_lines)
            changed_hexagram = self.find_hexagram(changed_pattern)
            
            if changed_hexagram:
                number, name = changed_hexagram
                self.current_changed = changed_hexagram
                self.changed_name_label.config(text=f"{number}번째 괘: {name}")
                self.load_interpretation(name, self.changed_text)
            else:
                self.current_changed = None
                self.changed_name_label.config(text="지괘를 찾을 수 없습니다")
                self.changed_text.delete(1.0, tk.END)
            
            # 최종 해석 계산
            self.calculate_final_interpretation(original_hexagram, changed_hexagram, moving_lines, lines)
        else:
            # 동효가 없으면 지괘 없음
            self.current_changed = None
            self.changed_name_label.config(text="동효가 없습니다")
            self.changed_text.delete(1.0, tk.END)
            self.changed_text.insert(tk.END, "동효를 체크하면 지괘가 표시됩니다.")
            
            # 동효가 없으면 본괘가 최종 해석
            self.current_final = original_hexagram
            self.current_rule = "동효 없음 → 본괘 전체 해석"
            self.final_name_label.config(text=f"최종: {original_hexagram[1]}")
            self.rule_label.config(text=self.current_rule)
            self.load_interpretation_with_highlight(original_hexagram[1], self.final_text, [])
    
    def calculate_final_interpretation(self, original_hexagram, changed_hexagram, moving_lines, lines):
        """동효 개수에 따른 최종 해석 계산"""
        moving_count = len(moving_lines)
        highlight_lines = []
        
        if moving_count == 1:
            # 1개: 본괘, 동효 칼라표시
            final_hexagram = original_hexagram
            highlight_lines = moving_lines
            rule = f"동효 1개 → 본괘 해석, {moving_lines[0]+1}효 강조"
            
        elif moving_count == 2:
            # 2개: 본괘, 큰 번호 효 칼라표시
            final_hexagram = original_hexagram
            highlight_lines = [max(moving_lines)]
            rule = f"동효 2개 → 본괘 해석, {max(moving_lines)+1}효 강조"
            
        elif moving_count == 3:
            # 3개: 하괘(1-3효)는 본괘, 상괘(4-6효)는 지괘의 조합으로 새로운 괘 생성
            # 지괘 계산 (모든 동효 반전)
            changed_lines = lines.copy()
            for i in moving_lines:
                changed_lines[i] = '2' if changed_lines[i] == '1' else '1'
            
            # 최종괘 생성: 하괘는 본괘, 상괘는 지괘
            final_lines = []
            for i in range(6):
                if i < 3:  # 하괘 (1-3효): 본괘 사용
                    final_lines.append(lines[i])
                else:  # 상괘 (4-6효): 지괘 사용
                    final_lines.append(changed_lines[i])
            
            # 새로운 괘 찾기
            final_pattern = ''.join(final_lines)
            final_hexagram = self.find_hexagram(final_pattern)
            
            if not final_hexagram:
                # 찾을 수 없으면 본괘 사용 (fallback)
                final_hexagram = original_hexagram
                rule = f"동효 3개 → 조합괘를 찾을 수 없어 본괘 해석"
            else:
                rule = f"동효 3개 → 하괘(본괘) + 상괘(지괘) 조합 해석"
            highlight_lines = []
            
        elif moving_count == 4:
            # 4개: 지괘, 동효가 아닌 것 중 작은 번호 칼라표시
            final_hexagram = changed_hexagram
            non_moving = [i for i in range(6) if i not in moving_lines]
            highlight_lines = [min(non_moving)] if non_moving else []
            rule = f"동효 4개 → 지괘 해석, {min(non_moving)+1 if non_moving else ''}효 강조"
            
        elif moving_count == 5:
            # 5개: 지괘, 동효가 아닌 것 칼라표시
            final_hexagram = changed_hexagram
            non_moving = [i for i in range(6) if i not in moving_lines]
            highlight_lines = non_moving
            rule = f"동효 5개 → 지괘 해석, {non_moving[0]+1 if non_moving else ''}효 강조"
            
        elif moving_count == 6:
            # 6개: 지괘 전체
            final_hexagram = changed_hexagram
            highlight_lines = []
            rule = f"동효 6개 → 지괘 전체 해석"
        
        # 현재 최종 결과 저장
        self.current_final = final_hexagram
        self.current_rule = rule
        self.current_highlight_lines = highlight_lines
        
        # 최종 해석 표시
        self.final_name_label.config(text=f"최종: {final_hexagram[1]}")
        self.rule_label.config(text=rule)
        self.load_interpretation_with_highlight(final_hexagram[1], self.final_text, highlight_lines)
    
    def sanitize_filename(self, filename):
        """파일명에서 사용할 수 없는 문자들을 제거하거나 변환"""
        # Windows/Mac/Linux에서 파일명에 사용할 수 없는 문자들
        invalid_chars = r'[<>:"/\\|?*]'
        # 특수문자를 언더스코어로 변환
        sanitized = re.sub(invalid_chars, '_', filename)
        # 연속된 공백이나 언더스코어를 하나로 줄임
        sanitized = re.sub(r'[_\s]+', '_', sanitized)
        # 시작과 끝의 공백, 언더스코어 제거
        sanitized = sanitized.strip('_. ')
        # 파일명이 너무 길면 자르기 (확장자 제외하고 100자로 제한)
        if len(sanitized) > 100:
            sanitized = sanitized[:100]
        # 빈 문자열이면 기본값 사용
        if not sanitized:
            sanitized = "점괘결과"
        return sanitized

    def save_result(self):
        """점괘 결과를 워드 파일로 저장"""
        if not self.current_original:
            messagebox.showwarning("저장 오류", "먼저 괘를 식별해주세요.")
            return
        
        # 출력 폴더 생성
        output_dir = Path("출력")
        output_dir.mkdir(exist_ok=True)
        
        # 질문 텍스트 가져오기
        question_text = self.current_question.strip() if self.current_question else ""
        if not question_text:
            question_text = "질문없음"
        
        # 파일명 생성 (날짜 + 질문)
        now = datetime.now()
        date_str = now.strftime('%Y%m%d_%H%M%S')
        
        # 질문을 파일명으로 사용할 수 있도록 처리
        sanitized_question = self.sanitize_filename(question_text)
        
        # 파일명 조합: 날짜_질문.docx
        filename = f"{date_str}_{sanitized_question}.docx"
        filepath = output_dir / filename
        
        try:
            # 새 워드 문서 생성
            doc = Document()
            
            # 제목
            doc.add_heading('주역 점괘 결과', 0)
            
            # 점괘 일시
            doc.add_paragraph(f"점괘 일시: {now.strftime('%Y년 %m월 %d일 %H시 %M분')}")
            doc.add_paragraph("")
            
            # 질문
            doc.add_heading('질문', level=1)
            question_text = self.current_question if self.current_question else "질문이 입력되지 않았습니다."
            doc.add_paragraph(question_text)
            doc.add_paragraph("")
            
            # 효 정보
            doc.add_heading('효 정보', level=1)
            effects_table = doc.add_table(rows=3, cols=7)
            effects_table.style = 'Table Grid'
            
            # 테이블 헤더
            header_cells = effects_table.rows[0].cells
            header_cells[0].text = "구분"
            for i in range(6):
                header_cells[i+1].text = f"{i+1}효"
            
            # 음양 정보
            yin_yang_cells = effects_table.rows[1].cells
            yin_yang_cells[0].text = "음양"
            for i, line in enumerate(self.current_lines):
                yin_yang_cells[i+1].text = "양" if line == "1" else "음"
            
            # 동효 정보
            moving_cells = effects_table.rows[2].cells
            moving_cells[0].text = "동효"
            for i in range(6):
                moving_cells[i+1].text = "○" if i in self.current_moving_lines else ""
            
            doc.add_paragraph("")
            
            # 본괘 정보
            doc.add_heading('본괘', level=1)
            if self.current_original:
                doc.add_paragraph(f"{self.current_original[0]}번째 괘: {self.current_original[1]}")
                doc.add_paragraph(f"효 패턴: {''.join(self.current_lines)}")
                
                # 본괘 해석 추가
                original_interpretation = self.get_interpretation_text(self.current_original[1])
                if original_interpretation:
                    doc.add_paragraph(original_interpretation)
            
            doc.add_paragraph("")
            
            # 지괘 정보 (동효가 있는 경우)
            if self.current_changed:
                doc.add_heading('지괘 (동효 반전)', level=1)
                doc.add_paragraph(f"{self.current_changed[0]}번째 괘: {self.current_changed[1]}")
                
                # 변경된 효 패턴 계산
                changed_lines = self.current_lines.copy()
                for i in self.current_moving_lines:
                    changed_lines[i] = '2' if changed_lines[i] == '1' else '1'
                doc.add_paragraph(f"효 패턴: {''.join(changed_lines)}")
                
                # 지괘 해석 추가
                changed_interpretation = self.get_interpretation_text(self.current_changed[1])
                if changed_interpretation:
                    doc.add_paragraph(changed_interpretation)
                
                doc.add_paragraph("")
            
            # 최종 해석 (강조 표시 포함)
            doc.add_heading('최종 해석', level=1)
            if self.current_final:
                doc.add_paragraph(f"최종 괘: {self.current_final[0]}번째 괘 - {self.current_final[1]}")
                doc.add_paragraph(f"해석 규칙: {self.current_rule}")
                
                # 최종 해석 추가 (강조 표시 포함)
                self.add_interpretation_with_highlight(doc, self.current_final[1], self.current_highlight_lines)
            
            # 파일 저장
            doc.save(filepath)
            
            messagebox.showinfo("저장 완료", f"점괘 결과가 저장되었습니다.\n파일 위치: {filepath}")
            
            # 출력 폴더를 자동으로 여는 기능 추가
            self.open_file_location(filepath)
            
        except Exception as e:
            messagebox.showerror("저장 오류", f"파일 저장 중 오류가 발생했습니다:\n{str(e)}")
    
    def open_file_location(self, filepath):
        """파일 위치를 운영체제별로 열기"""
        try:
            if platform.system() == "Windows":
                # Windows: 파일을 선택한 상태로 탐색기 열기
                subprocess.run(['explorer', '/select,', str(filepath)])
            elif platform.system() == "Darwin":  # macOS
                # macOS: Finder에서 파일을 선택한 상태로 열기
                subprocess.run(['open', '-R', str(filepath)])
            else:  # Linux
                # Linux: 파일 매니저에서 폴더 열기
                folder = filepath.parent
                subprocess.run(['xdg-open', str(folder)])
        except Exception as e:
            # 실패할 경우 폴더만 열기 시도
            try:
                folder = filepath.parent
                if platform.system() == "Windows":
                    os.startfile(str(folder))
                elif platform.system() == "Darwin":
                    subprocess.run(['open', str(folder)])
                else:
                    subprocess.run(['xdg-open', str(folder)])
            except Exception as e2:
                print(f"폴더 열기 실패: {e2}")
    
    def add_interpretation_with_highlight(self, doc, hexagram_name, highlight_lines):
        """워드 문서에 해석을 추가하면서 특정 효를 강조 표시"""
        doc_path = Path(f"해석/{hexagram_name}.docx")
        
        if not doc_path.exists():
            doc.add_paragraph(f"'{hexagram_name}' 해석 파일을 찾을 수 없습니다.")
            return
        
        try:
            source_doc = Document(doc_path)
            
            for paragraph in source_doc.paragraphs:
                if paragraph.text.strip():
                    # 새 단락 추가
                    new_paragraph = doc.add_paragraph()
                    
                    # 강조할 효가 있는지 확인
                    should_highlight = False
                    for highlight_line in highlight_lines:
                        effect_num = highlight_line + 1
                        if paragraph.text.strip().startswith(f"{effect_num}효"):
                            should_highlight = True
                            break
                    
                    if should_highlight:
                        # 강조 표시 (빨간색, 볼드)
                        run = new_paragraph.add_run(paragraph.text)
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 0, 0)  # 빨간색
                        run.font.size = Pt(14)  # 크기 증가
                    else:
                        # 일반 텍스트
                        new_paragraph.add_run(paragraph.text)
                        
        except Exception as e:
            doc.add_paragraph(f"해석 파일을 읽는 중 오류가 발생했습니다: {str(e)}")
    
    def find_hexagram(self, pattern):
        for number, (name, binary_pattern) in self.hexagrams.items():
            if pattern == binary_pattern:
                return (number, name)
        return None
    
    def load_interpretation(self, hexagram_name, text_widget):
        # 워드 파일 경로
        doc_path = Path(f"해석/{hexagram_name}.docx")
        
        if not doc_path.exists():
            text_widget.delete(1.0, tk.END)
            text_widget.insert(tk.END, 
                f"'{hexagram_name}.docx' 파일을 찾을 수 없습니다.\n\n"
                f"해석 폴더에 해당 파일을 생성해주세요.\n"
                f"파일 경로: {doc_path.absolute()}")
            return
        
        try:
            # 워드 파일 읽기
            doc = Document(doc_path)
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # 텍스트 영역에 표시 (전체 설명과 효 설명들 사이, 효들 사이, 효와 전체 사이에 빈 줄 추가)
            text_widget.delete(1.0, tk.END)
            if content:
                formatted_content = []
                
                for i, line in enumerate(content):
                    # 현재 줄이 효 설명인지 확인 (1효, 2효, 3효, 4효, 5효, 6효로 시작하는지)
                    current_is_effect = any(line.strip().startswith(f"{j}효") for j in range(1, 7))
                    
                    # 이전 줄이 있는 경우 간격 처리
                    if i > 0:
                        prev_line = content[i-1]
                        prev_is_effect = any(prev_line.strip().startswith(f"{j}효") for j in range(1, 7))
                        
                        # 경우 1: 이전 줄이 효가 아니고 현재 줄이 효 (전체 설명과 첫 번째 효 사이)
                        # 경우 2: 이전 줄이 효이고 현재 줄도 효 (효와 효 사이)
                        # 경우 3: 이전 줄이 효이고 현재 줄이 효가 아님 (효와 전체 설명 사이)
                        if (not prev_is_effect and current_is_effect) or \
                           (prev_is_effect and current_is_effect) or \
                           (prev_is_effect and not current_is_effect):
                            formatted_content.append("")
                    
                    formatted_content.append(line)
                
                text_widget.insert(tk.END, '\n'.join(formatted_content))
            else:
                text_widget.insert(tk.END, "파일이 비어있습니다.")
                
        except Exception as e:
            text_widget.delete(1.0, tk.END)
            text_widget.insert(tk.END, f"파일을 읽는 중 오류가 발생했습니다:\n{str(e)}")
    
    def clear_inputs(self):
        # 질문 지우기
        self.question_entry.delete(0, tk.END)
        
        # 효 입력 지우기
        for var in self.line_vars:
            var.set("1")
        for var in self.moving_vars:
            var.set(False)
            
        # 결과 지우기
        self.original_name_label.config(text="")
        self.changed_name_label.config(text="")
        self.final_name_label.config(text="")
        self.rule_label.config(text="")
        self.original_text.delete(1.0, tk.END)
        self.changed_text.delete(1.0, tk.END)
        self.final_text.delete(1.0, tk.END)
        
        # 현재 데이터 초기화
        self.current_question = ""
        self.current_lines = []
        self.current_moving_lines = []
        self.current_original = None
        self.current_changed = None
        self.current_final = None
        self.current_rule = ""
        self.current_highlight_lines = []  # 강조할 효 번호들
        
        # 질문 입력칸에 포커스
        self.question_entry.focus()
    
    def get_interpretation_text(self, hexagram_name):
        """괘 해석 텍스트를 파일에서 읽어오기 (전체 설명과 효 설명들 사이, 효들 사이, 효와 전체 사이에 빈 줄 추가)"""
        doc_path = Path(f"해석/{hexagram_name}.docx")
        
        if not doc_path.exists():
            return f"'{hexagram_name}' 해석 파일을 찾을 수 없습니다."
        
        try:
            doc = Document(doc_path)
            content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            if content:
                # 전체 설명과 효 설명들 사이, 효들 사이, 효와 전체 사이에 빈 줄 추가
                formatted_content = []
                
                for i, line in enumerate(content):
                    # 현재 줄이 효 설명인지 확인
                    current_is_effect = any(line.strip().startswith(f"{j}효") for j in range(1, 7))
                    
                    # 이전 줄이 있는 경우 간격 처리
                    if i > 0:
                        prev_line = content[i-1]
                        prev_is_effect = any(prev_line.strip().startswith(f"{j}효") for j in range(1, 7))
                        
                        # 경우 1: 이전 줄이 효가 아니고 현재 줄이 효 (전체 설명과 첫 번째 효 사이)
                        # 경우 2: 이전 줄이 효이고 현재 줄도 효 (효와 효 사이)
                        # 경우 3: 이전 줄이 효이고 현재 줄이 효가 아님 (효와 전체 설명 사이)
                        if (not prev_is_effect and current_is_effect) or \
                           (prev_is_effect and current_is_effect) or \
                           (prev_is_effect and not current_is_effect):
                            formatted_content.append("")
                    
                    formatted_content.append(line)
                
                return '\n'.join(formatted_content)
            else:
                return "해석 내용이 없습니다."
        except Exception as e:
            return f"해석 파일을 읽는 중 오류가 발생했습니다: {str(e)}"
    
    def load_interpretation_with_highlight(self, hexagram_name, text_widget, highlight_lines):
        """해석을 로드하고 특정 효를 칼라 표시 (효 사이에 빈 줄 추가)"""
        # 기본 해석 로드 (이미 효 사이에 빈 줄이 추가됨)
        self.load_interpretation(hexagram_name, text_widget)
        
        # 칼라 표시할 효가 있으면 하이라이트
        if highlight_lines:
            text_widget.tag_configure("highlight", 
                background="#FFFF00", foreground="#FF0000", 
                font=('맑은 고딕', self.font_sizes['text'], 'bold'))
            
            content = text_widget.get(1.0, tk.END)
            lines_list = content.split('\n')
            
            for highlight_line in highlight_lines:
                effect_num = highlight_line + 1
                # 해당 효를 찾아서 하이라이트
                for i, line in enumerate(lines_list):
                    if line.strip().startswith(f"{effect_num}효"):
                        # 해당 라인의 시작과 끝 위치 계산
                        start_pos = sum(len(l) + 1 for l in lines_list[:i])
                        end_pos = start_pos + len(line)
                        
                        # 하이라이트 적용
                        text_widget.tag_add("highlight", f"1.0+{start_pos}c", f"1.0+{end_pos}c")
                        break

if __name__ == "__main__":
    root = tk.Tk()
    app = IChingReader(root)
    root.mainloop() 