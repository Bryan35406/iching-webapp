from flask import Flask, render_template, request, jsonify, send_file, make_response
import os
import random
from pathlib import Path
import docx
from docx import Document
from docx.shared import RGBColor, Pt
from datetime import datetime
import zipfile
import io
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import sqlite3
import json
import uuid
import tempfile
import re
import platform
import sys

# 이메일 설정 import
try:
    from email_config import GMAIL_APP_PASSWORD, SENDER_EMAIL, RECIPIENT_EMAIL
except ImportError:
    # 설정 파일이 없으면 기본값 사용
    GMAIL_APP_PASSWORD = "your-app-password"
    SENDER_EMAIL = "your-email@gmail.com"
    RECIPIENT_EMAIL = "recipient@gmail.com"

app = Flask(__name__)

# Vercel 환경에서 임시 디렉토리 사용
def get_temp_dir():
    """임시 디렉토리 경로 반환"""
    if os.environ.get('VERCEL'):
        return Path('/tmp')
    else:
        return Path('.')

def get_output_dir():
    """출력 디렉토리 경로 반환"""
    if os.environ.get('VERCEL'):
        # Vercel 환경에서는 /tmp 디렉토리만 쓰기 가능
        return Path('/tmp')
    else:
        # 로컬 환경에서는 출력 폴더 생성
        base_dir = Path('.')
        output_dir = base_dir / "출력"
        output_dir.mkdir(exist_ok=True)
        return output_dir

def get_db_path():
    """데이터베이스 파일 경로 반환"""
    base_dir = get_temp_dir()
    return base_dir / 'iching_history.db'

class IChingWeb:
    def __init__(self):
        # 64괘 정보 (괘번호: (괘이름, 이진코드)) - 정확한 전통 패턴
        self.hexagrams = {
            1: ("1괘 - 건위천", "111111"),      # 건위천(乾爲天)
            2: ("2괘 - 곤위지", "222222"),      # 곤위지(坤爲地)
            3: ("3괘 - 수뢰둔", "122212"),      # 수뢰둔(水雷屯)
            4: ("4괘 - 산수몽", "212221"),      # 산수몽(山水蒙)
            5: ("5괘 - 수천수", "111212"),      # 수천수(水天需)
            6: ("6괘 - 천수송", "212111"),      # 천수송(天水訟)
            7: ("7괘 - 지수사", "212222"),      # 지수사(地水師)
            8: ("8괘 - 수지비", "222212"),      # 수지비(水地比)
            9: ("9괘 - 풍천소축", "111211"),    # 풍천소축(風天小畜)
            10: ("10괘 - 천택리", "112111"),    # 천택리(天澤履)
            11: ("11괘 - 지천태", "111222"),    # 지천태(地天泰)
            12: ("12괘 - 천지비", "222111"),    # 천지비(天地否)
            13: ("13괘 - 천화동인", "121111"),  # 천화동인(天火同人)
            14: ("14괘 - 화천대유", "111121"),  # 화천대유(火天大有)
            15: ("15괘 - 지산겸", "221222"),    # 지산겸(地山謙)
            16: ("16괘 - 뢰지예", "222122"),    # 뢰지예(雷地豫)
            17: ("17괘 - 택뢰수", "122112"),    # 택뢰수(澤雷隨)
            18: ("18괘 - 산풍고", "211221"),    # 산풍고(山風蠱)
            19: ("19괘 - 지택림", "112222"),    # 지택림(地澤臨)
            20: ("20괘 - 풍지관", "222211"),    # 풍지관(風地觀)
            21: ("21괘 - 화뢰서합", "122121"),  # 화뢰서합(火雷噬嗑)
            22: ("22괘 - 산화비", "121221"),    # 산화비(山火賁)
            23: ("23괘 - 산지박", "222221"),    # 산지박(山地剝)
            24: ("24괘 - 지뢰복", "122222"),    # 지뢰복(地雷復)
            25: ("25괘 - 천뢰무망", "122111"),  # 천뢰무망(天雷无妄)
            26: ("26괘 - 산천대축", "111221"),  # 산천대축(山天大畜)
            27: ("27괘 - 산뢰이", "122221"),    # 산뢰이(山雷頤)
            28: ("28괘 - 택풍대과", "211112"),  # 택풍대과(澤風大過)
            29: ("29괘 - 감위수", "212212"),    # 감위수(坎爲水)
            30: ("30괘 - 리위화", "121121"),    # 리위화(離爲火)
            31: ("31괘 - 택산함", "221112"),    # 택산함(澤山咸)
            32: ("32괘 - 뢰풍항", "211122"),    # 뢰풍항(雷風恆)
            33: ("33괘 - 천산둔", "221111"),    # 천산둔(天山遯)
            34: ("34괘 - 뢰천대장", "111122"),  # 뢰천대장(雷天大壯)
            35: ("35괘 - 화지진", "222121"),    # 화지진(火地晉)
            36: ("36괘 - 지화명이", "121222"),  # 지화명이(地火明夷)
            37: ("37괘 - 풍화가인", "121211"),  # 풍화가인(風火家人)
            38: ("38괘 - 화택규", "112121"),    # 화택규(火澤睽)
            39: ("39괘 - 수산건", "221212"),    # 수산건(Water Mountain Obstruction)
            40: ("40괘 - 뢰수해", "212122"),    # 뢰수해(Thunder Water Deliverance)
            41: ("41괘 - 산택손", "112221"),    # 산택손(Mountain Lake Decrease)
            42: ("42괘 - 풍뢰익", "122211"),    # 풍뢰익(Wind Thunder Increase)
            43: ("43괘 - 택천쾌", "111112"),    # 택천쾌(Lake Heaven Breakthrough)
            44: ("44괘 - 천풍구", "211111"),    # 천풍구(Heaven Wind Coming-to-meet)
            45: ("45괘 - 택지취", "222112"),    # 택지취(Lake Earth Gathering)
            46: ("46괘 - 지풍승", "211222"),    # 지풍승(Earth Wind Pushing Upward)
            47: ("47괘 - 택수곤", "212112"),    # 택수곤(Lake Water Oppression)
            48: ("48괘 - 수풍정", "211212"),    # 수풍정(Water Wind The Well)
            49: ("49괘 - 택화혁", "121112"),    # 택화혁(Lake Fire Revolution)
            50: ("50괘 - 화풍정", "211121"),    # 화풍정(Fire Wind The Cauldron)
            51: ("51괘 - 진위뢰", "122122"),    # 진위뢰(震爲雷)
            52: ("52괘 - 간위산", "221221"),    # 간위산(艮爲山)
            53: ("53괘 - 풍산점", "221211"),    # 풍산점(Wind Mountain Development)
            54: ("54괘 - 뢰택귀매", "112122"),  # 뢰택귀매(Thunder Lake Marrying Maiden)
            55: ("55괘 - 뢰화풍", "121122"),    # 뢰화풍(Thunder Fire Abundance)
            56: ("56괘 - 화산려", "221121"),    # 화산려(Fire Mountain The Wanderer)
            57: ("57괘 - 손위풍", "211211"),    # 손위풍(巽爲風)
            58: ("58괘 - 태위택", "112112"),    # 태위택(兌爲澤)
            59: ("59괘 - 풍수환", "212211"),    # 풍수환(Wind Water Dispersion)
            60: ("60괘 - 수택절", "112212"),    # 수택절(Water Lake Limitation)
            61: ("61괘 - 풍택중부", "112211"),  # 풍택중부(Wind Lake Inner Truth)
            62: ("62괘 - 뢰산소과", "221122"),  # 뢰산소과(Thunder Mountain Small Exceeding)
            63: ("63괘 - 수화기제", "121212"),  # 수화기제(Water Fire After Completion)
            64: ("64괘 - 화수미제", "212121"),  # 화수미제(Fire Water Before Completion)
        }
    
    def find_hexagram(self, pattern):
        for number, (name, binary_pattern) in self.hexagrams.items():
            if pattern == binary_pattern:
                return (number, name)
        return None
    
    def extract_hexagram_number(self, hexagram_name):
        """괘 이름에서 괘 번호만 추출 (예: '3괘 - 수뢰둔' -> '3괘')"""
        if ' - ' in hexagram_name:
            return hexagram_name.split(' - ')[0]
        return hexagram_name
    
    def get_interpretation_text(self, hexagram_name):
        """해석 파일에서 텍스트 읽기"""
        try:
            # 괘 번호만 추출하여 파일명 생성
            file_name = self.extract_hexagram_number(hexagram_name)
            
            # PyInstaller 환경 확인
            if getattr(sys, 'frozen', False):
                # PyInstaller로 빌드된 실행 파일인 경우
                if hasattr(sys, '_MEIPASS'):
                    # 임시 폴더에서 실행 중인 경우
                    application_path = Path(sys._MEIPASS)
                    print(f"DEBUG: PyInstaller 임시 경로: {application_path}")
                else:
                    # 실행 파일 위치
                    application_path = Path(sys.executable).parent
                    print(f"DEBUG: PyInstaller 실행 파일 경로: {application_path}")
            else:
                # 일반 Python 스크립트로 실행 중인 경우
                application_path = Path(__file__).parent
                print(f"DEBUG: Python 스크립트 경로: {application_path}")
            
            # 현재 실행 경로 확인
            current_dir = Path.cwd()
            print(f"DEBUG: 현재 작업 디렉토리: {current_dir}")
            
            # Vercel 환경에서는 해석 파일을 프로젝트 루트에서 찾기
            if os.environ.get('VERCEL'):
                interpretation_path = Path(f"해석/{file_name}.docx")
            else:
                # 로컬 환경에서 여러 경로 시도 (PyInstaller 환경 및 macOS .app 번들 포함)
                possible_paths = [
                    # PyInstaller 환경 우선 경로
                    application_path / "해석" / f"{file_name}.docx",
                    current_dir / "해석" / f"{file_name}.docx",
                    
                    # 복사된 해석 폴더들
                    Path("/Users/bpark/Desktop/0_Python/주역/dist/해석") / f"{file_name}.docx",
                    Path("/Users/bpark/Desktop/0_Python/주역/dist/주역괘해석/해석") / f"{file_name}.docx",
                    Path("/Users/bpark/Desktop/0_Python/주역/dist/주역괘해석.app/Contents/MacOS/해석") / f"{file_name}.docx",
                    
                    # 기본 경로들
                    Path(f"해석/{file_name}.docx"),  # 현재 디렉토리에서
                    Path(f"../해석/{file_name}.docx"),  # 상위 디렉토리에서
                    Path(f"../../해석/{file_name}.docx"),  # 두 단계 상위 디렉토리에서
                    Path(f"/Users/bpark/Desktop/0_Python/주역/해석/{file_name}.docx"),  # 원본 절대 경로
                    
                    # PyInstaller _internal 폴더를 고려한 경로
                    Path(current_dir.parent.parent / "해석" / f"{file_name}.docx"),
                    Path(current_dir.parent / "해석" / f"{file_name}.docx"),
                ]
                
                print(f"DEBUG: 시도할 경로들:")
                for i, path in enumerate(possible_paths):
                    exists = path.exists()
                    print(f"  {i+1}. {path} -> 존재: {exists}")
                    if exists:
                        print(f"      파일 크기: {path.stat().st_size} bytes")
                
                interpretation_path = None
                for path in possible_paths:
                    if path.exists():
                        interpretation_path = path
                        print(f"DEBUG: 찾은 경로: {interpretation_path}")
                        break
                
                if interpretation_path is None:
                    # 해석 폴더 자체를 찾기 위한 추가 시도
                    base_paths = [
                        application_path,
                        Path("/Users/bpark/Desktop/0_Python/주역"),
                        Path("/Users/bpark/Desktop/0_Python/주역/dist"),
                        Path("/Users/bpark/Desktop/0_Python/주역/dist/주역괘해석"),
                        Path("/Users/bpark/Desktop/0_Python/주역/dist/주역괘해석.app/Contents/MacOS"),
                        current_dir,
                        current_dir.parent,
                        current_dir.parent.parent,
                        current_dir.parent.parent.parent,
                    ]
                    
                    print(f"DEBUG: 해석 폴더 검색...")
                    for base_path in base_paths:
                        interpretation_dir = base_path / "해석"
                        target_file = interpretation_dir / f"{file_name}.docx"
                        exists = target_file.exists()
                        print(f"  검색: {target_file} -> 존재: {exists}")
                        if exists:
                            interpretation_path = target_file
                            print(f"DEBUG: 최종 찾은 경로: {interpretation_path}")
                            break
                    
                    if interpretation_path is None:
                        error_msg = f"{file_name}에 대한 해석 파일을 찾을 수 없습니다.\n"
                        error_msg += f"현재 작업 디렉토리: {current_dir}\n"
                        error_msg += f"애플리케이션 경로: {application_path}\n"
                        error_msg += f"실행 파일 경로: {sys.executable if getattr(sys, 'frozen', False) else 'N/A'}\n"
                        error_msg += f"찾는 파일: {file_name}.docx\n"
                        error_msg += "시도한 모든 경로에서 파일을 찾을 수 없습니다."
                        return error_msg
            
            if not interpretation_path.exists():
                return f"{file_name}에 대한 해석 파일을 찾을 수 없습니다."
            
            doc = Document(interpretation_path)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            return '\n'.join(text_content) if text_content else f"{file_name}에 대한 해석 내용이 없습니다."
            
        except Exception as e:
            file_name = self.extract_hexagram_number(hexagram_name)
            return f"{file_name} 해석 파일 읽기 오류: {str(e)}\n디버깅 정보: 현재 디렉토리={Path.cwd()}, 실행파일={getattr(sys, 'executable', 'N/A')}"
    
    def calculate_final_interpretation(self, original_hexagram, changed_hexagram, moving_lines, lines):
        """동효 개수에 따른 최종 해석 계산"""
        moving_count = len(moving_lines)
        highlight_lines = []
        
        if moving_count == 1:
            final_hexagram = original_hexagram
            highlight_lines = moving_lines
            rule = f"동효 1개 → 본괘 해석, {moving_lines[0]+1}효 강조"
            
        elif moving_count == 2:
            final_hexagram = original_hexagram
            highlight_lines = [max(moving_lines)]
            rule = f"동효 2개 → 본괘 해석, {max(moving_lines)+1}효 강조"
            
        elif moving_count == 3:
            # 3개: 하괘(1-3효)는 본괘, 상괘(4-6효)는 지괘의 조합으로 새로운 괘 생성
            # 지괘 계산 (모든 동효 반전)
            changed_lines = lines.copy()
            for i in moving_lines:
                changed_lines[i] = '2' if changed_lines[i] == '1' else '1'
            
            # 최종괘 생성: 하괘는 본괘, 상괘는 지괘
            final_lines = []
            for i in range(6):
                if i < 3:  # 하괘 (1-3효): 본괘 사용
                    final_lines.append(lines[i])
                else:  # 상괘 (4-6효): 지괘 사용
                    final_lines.append(changed_lines[i])
            
            # 새로운 괘 찾기
            final_pattern = ''.join(final_lines)
            final_hexagram = self.find_hexagram(final_pattern)
            
            if not final_hexagram:
                # 찾을 수 없으면 본괘 사용 (fallback)
                final_hexagram = original_hexagram
                rule = f"동효 3개 → 조합괘를 찾을 수 없어 본괘 해석"
            else:
                rule = f"동효 3개 → 하괘(본괘) + 상괘(지괘) 조합 해석"
            highlight_lines = []
            
        elif moving_count == 4:
            final_hexagram = changed_hexagram
            non_moving = [i for i in range(6) if i not in moving_lines]
            highlight_lines = [min(non_moving)] if non_moving else []
            rule = f"동효 4개 → 지괘 해석, {min(non_moving)+1 if non_moving else ''}효 강조"
            
        elif moving_count == 5:
            final_hexagram = changed_hexagram
            non_moving = [i for i in range(6) if i not in moving_lines]
            highlight_lines = non_moving
            rule = f"동효 5개 → 지괘 해석, {non_moving[0]+1 if non_moving else ''}효 강조"
            
        elif moving_count == 6:
            final_hexagram = changed_hexagram
            highlight_lines = []
            rule = f"동효 6개 → 지괘 전체 해석"
        
        return final_hexagram, rule, highlight_lines

# 글로벌 인스턴스
iching = IChingWeb()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/identify', methods=['POST'])
def identify_hexagram():
    try:
        data = request.json
        question = data.get('question', '')
        lines = data.get('lines', [])
        moving_lines = data.get('moving_lines', [])
        
        # 입력 검증
        if len(lines) != 6:
            return jsonify({'error': '6개의 효를 모두 입력해주세요.'}), 400
        
        for i, line in enumerate(lines):
            if line not in ['1', '2']:
                return jsonify({'error': f'{i+1}효는 1(양) 또는 2(음)만 입력 가능합니다.'}), 400
        
        # 본괘 찾기
        original_pattern = ''.join(lines)
        original_hexagram = iching.find_hexagram(original_pattern)
        
        if not original_hexagram:
            return jsonify({'error': '해당하는 본괘를 찾을 수 없습니다.'}), 400
        
        result = {
            'original': {
                'number': original_hexagram[0],
                'name': original_hexagram[1],
                'interpretation': iching.get_interpretation_text(original_hexagram[1])
            }
        }
        
        # 동효가 있으면 지괘 계산
        if moving_lines:
            changed_lines = lines.copy()
            for i in moving_lines:
                changed_lines[i] = '2' if changed_lines[i] == '1' else '1'
            
            changed_pattern = ''.join(changed_lines)
            changed_hexagram = iching.find_hexagram(changed_pattern)
            
            if changed_hexagram:
                result['changed'] = {
                    'number': changed_hexagram[0],
                    'name': changed_hexagram[1],
                    'interpretation': iching.get_interpretation_text(changed_hexagram[1])
                }
                
                # 최종 해석 계산
                final_hexagram, rule, highlight_lines = iching.calculate_final_interpretation(
                    original_hexagram, changed_hexagram, moving_lines, lines
                )
                
                result['final'] = {
                    'number': final_hexagram[0],
                    'name': final_hexagram[1],
                    'interpretation': iching.get_interpretation_text(final_hexagram[1]),
                    'rule': rule,
                    'highlight_lines': highlight_lines
                }
        else:
            # 동효가 없으면 본괘가 최종 해석
            result['final'] = {
                'number': original_hexagram[0],
                'name': original_hexagram[1],
                'interpretation': iching.get_interpretation_text(original_hexagram[1]),
                'rule': '동효 없음 → 본괘 전체 해석',
                'highlight_lines': []
            }
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'서버 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/api/auto-generate', methods=['POST'])
def auto_generate():
    """자동으로 괘와 동효를 생성"""
    try:
        # 6개 효 랜덤 생성
        lines = [random.choice(['1', '2']) for _ in range(6)]
        
        # 동효 랜덤 생성 (확률 가중치 적용)
        moving_count = random.choices(
            population=[0, 1, 2, 3], 
            weights=[50, 30, 15, 5], 
            k=1
        )[0]
        
        moving_lines = []
        if moving_count > 0:
            moving_lines = random.sample(range(6), moving_count)
        
        return jsonify({
            'lines': lines,
            'moving_lines': moving_lines
        })
        
    except Exception as e:
        return jsonify({'error': f'자동 생성 중 오류가 발생했습니다: {str(e)}'}), 500

def sanitize_filename(filename):
    """파일명에서 사용할 수 없는 문자들을 제거하거나 변환"""
    # Windows/Mac/Linux에서 파일명에 사용할 수 없는 문자들
    invalid_chars = r'[<>:"/\\|?*]'
    # 특수문자를 언더스코어로 변환
    sanitized = re.sub(invalid_chars, '_', filename)
    # 연속된 공백이나 언더스코어를 하나로 줄임
    sanitized = re.sub(r'[_\s]+', '_', sanitized)
    # 시작과 끝의 공백, 언더스코어 제거
    sanitized = sanitized.strip('_. ')
    # 파일명이 너무 길면 자르기 (확장자 제외하고 100자로 제한)
    if len(sanitized) > 100:
        sanitized = sanitized[:100]
    # 빈 문자열이면 기본값 사용
    if not sanitized:
        sanitized = "점괘결과"
    return sanitized

@app.route('/api/save', methods=['POST'])
def save_result():
    """점괘 결과를 워드 파일로 저장하고 다운로드"""
    try:
        data = request.json
        
        # 출력 폴더 생성
        output_dir = get_output_dir()
        
        # 질문 텍스트 가져오기
        question_text = data.get('question', '').strip()
        if not question_text:
            question_text = "질문없음"
        
        # 파일명 생성 (날짜 + 질문)
        now = datetime.now()
        date_str = now.strftime('%Y%m%d_%H%M%S')
        
        # 질문을 파일명으로 사용할 수 있도록 처리
        sanitized_question = sanitize_filename(question_text)
        
        # 파일명 조합: 날짜_질문.docx
        filename = f"{date_str}_{sanitized_question}.docx"
        filepath = output_dir / filename
        
        # 워드 문서 생성
        doc = Document()
        
        # 제목
        doc.add_heading('주역 점괘 결과', 0)
        
        # 점괘 일시
        doc.add_paragraph(f"점괘 일시: {now.strftime('%Y년 %m월 %d일 %H시 %M분')}")
        doc.add_paragraph("")
        
        # 질문
        doc.add_heading('질문', level=1)
        doc.add_paragraph(question_text)
        doc.add_paragraph("")
        
        # 효 정보
        doc.add_heading('효 정보', level=1)
        lines = data.get('lines', [])
        moving_lines = data.get('moving_lines', [])
        
        effects_table = doc.add_table(rows=3, cols=7)
        effects_table.style = 'Table Grid'
        
        # 테이블 헤더
        header_cells = effects_table.rows[0].cells
        header_cells[0].text = "구분"
        for i in range(6):
            header_cells[i+1].text = f"{i+1}효"
        
        # 음양 정보
        yin_yang_cells = effects_table.rows[1].cells
        yin_yang_cells[0].text = "음양"
        for i, line in enumerate(lines):
            yin_yang_cells[i+1].text = "양" if line == "1" else "음"
        
        # 동효 정보
        moving_cells = effects_table.rows[2].cells
        moving_cells[0].text = "동효"
        for i in range(6):
            moving_cells[i+1].text = "○" if i in moving_lines else ""
        
        doc.add_paragraph("")
        
        # 결과 정보 추가
        original = data.get('original', {})
        changed = data.get('changed', {})
        final = data.get('final', {})
        
        # 본괘
        if original:
            doc.add_heading('본괘', level=1)
            doc.add_paragraph(f"{original['number']}번째 괘: {original['name']}")
            doc.add_paragraph(f"효 패턴: {''.join(lines)}")
            doc.add_paragraph(original.get('interpretation', ''))
            doc.add_paragraph("")
        
        # 지괘
        if changed:
            doc.add_heading('지괘 (동효 반전)', level=1)
            doc.add_paragraph(f"{changed['number']}번째 괘: {changed['name']}")
            
            changed_lines = lines.copy()
            for i in moving_lines:
                changed_lines[i] = '2' if changed_lines[i] == '1' else '1'
            doc.add_paragraph(f"효 패턴: {''.join(changed_lines)}")
            
            doc.add_paragraph(changed.get('interpretation', ''))
            doc.add_paragraph("")
        
        # 최종 해석
        if final:
            doc.add_heading('최종 해석', level=1)
            doc.add_paragraph(f"최종 괘: {final['number']}번째 괘 - {final['name']}")
            doc.add_paragraph(f"해석 규칙: {final.get('rule', '')}")
            
            # 최종 해석 추가 (강조 표시 포함)
            interpretation_lines = final.get('interpretation', '').split('\n')
            highlight_lines = final.get('highlight_lines', [])
            
            for line in interpretation_lines:
                if line.strip():
                    paragraph = doc.add_paragraph()
                    
                    # 강조할 효인지 확인
                    should_highlight = False
                    for highlight_line in highlight_lines:
                        effect_num = highlight_line + 1
                        if line.strip().startswith(f"{effect_num}효"):
                            should_highlight = True
                            break
                    
                    if should_highlight:
                        run = paragraph.add_run(line)
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.font.size = Pt(14)
                    else:
                        paragraph.add_run(line)
        
        # 파일 저장
        doc.save(filepath)
        
        return jsonify({
            'success': True,
            'filename': filename,
            'download_url': f'/api/download/{filename}',
            'message': f'점괘 결과가 저장되었습니다: {filename}'
        })
        
    except Exception as e:
        return jsonify({'error': f'저장 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/api/download/<filename>')
def download_file(filename):
    """워드 파일 다운로드"""
    try:
        # 보안을 위해 파일명 검증
        if not filename.endswith('.docx') or '..' in filename:
            return jsonify({'error': '잘못된 파일명입니다.'}), 400
        
        filepath = get_output_dir() / filename
        
        if not filepath.exists():
            return jsonify({'error': '파일을 찾을 수 없습니다.'}), 404
        
        # 파일을 바이너리로 읽어서 Response 생성
        try:
            with open(filepath, 'rb') as f:
                file_data = f.read()
            
            response = make_response(file_data)
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            # 안전한 ASCII 파일명 사용
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
            response.headers['Content-Length'] = len(file_data)
            
            return response
        except Exception as file_error:
            return jsonify({'error': f'파일 읽기 오류: {str(file_error)}'}), 500
        
    except Exception as e:
        return jsonify({'error': f'다운로드 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/api/send-email', methods=['POST'])
def send_email():
    """점괘 결과를 이메일로 전송"""
    try:
        data = request.json
        
        # 출력 폴더 생성
        output_dir = get_output_dir()
        
        # 질문 텍스트 가져오기
        question_text = data.get('question', '').strip()
        if not question_text:
            question_text = "질문없음"
        
        # 파일명 생성 (날짜 + 질문)
        now = datetime.now()
        date_str = now.strftime('%Y%m%d_%H%M%S')
        
        # 질문을 파일명으로 사용할 수 있도록 처리
        sanitized_question = sanitize_filename(question_text)
        
        # 파일명 조합: 날짜_질문.docx
        filename = f"{date_str}_{sanitized_question}.docx"
        filepath = output_dir / filename
        
        # 워드 문서 생성 (기존 save_result와 동일한 로직)
        doc = Document()
        
        # 제목
        doc.add_heading('주역 점괘 결과', 0)
        
        # 점괘 일시
        doc.add_paragraph(f"점괘 일시: {now.strftime('%Y년 %m월 %d일 %H시 %M분')}")
        doc.add_paragraph("")
        
        # 질문
        doc.add_heading('질문', level=1)
        doc.add_paragraph(question_text)
        doc.add_paragraph("")
        
        # 효 정보
        doc.add_heading('효 정보', level=1)
        lines = data.get('lines', [])
        moving_lines = data.get('moving_lines', [])
        
        effects_table = doc.add_table(rows=3, cols=7)
        effects_table.style = 'Table Grid'
        
        # 테이블 헤더
        header_cells = effects_table.rows[0].cells
        header_cells[0].text = "구분"
        for i in range(6):
            header_cells[i+1].text = f"{i+1}효"
        
        # 음양 정보
        yin_yang_cells = effects_table.rows[1].cells
        yin_yang_cells[0].text = "음양"
        for i, line in enumerate(lines):
            yin_yang_cells[i+1].text = "양" if line == "1" else "음"
        
        # 동효 정보
        moving_cells = effects_table.rows[2].cells
        moving_cells[0].text = "동효"
        for i in range(6):
            moving_cells[i+1].text = "○" if i in moving_lines else ""
        
        doc.add_paragraph("")
        
        # 결과 정보 추가
        original = data.get('original', {})
        changed = data.get('changed', {})
        final = data.get('final', {})
        
        # 본괘
        if original:
            doc.add_heading('본괘', level=1)
            doc.add_paragraph(f"{original['number']}번째 괘: {original['name']}")
            doc.add_paragraph(f"효 패턴: {''.join(lines)}")
            doc.add_paragraph(original.get('interpretation', ''))
            doc.add_paragraph("")
        
        # 지괘
        if changed:
            doc.add_heading('지괘 (동효 반전)', level=1)
            doc.add_paragraph(f"{changed['number']}번째 괘: {changed['name']}")
            
            changed_lines = lines.copy()
            for i in moving_lines:
                changed_lines[i] = '2' if changed_lines[i] == '1' else '1'
            doc.add_paragraph(f"효 패턴: {''.join(changed_lines)}")
            
            doc.add_paragraph(changed.get('interpretation', ''))
            doc.add_paragraph("")
        
        # 최종 해석
        if final:
            doc.add_heading('최종 해석', level=1)
            doc.add_paragraph(f"최종 괘: {final['number']}번째 괘 - {final['name']}")
            doc.add_paragraph(f"해석 규칙: {final.get('rule', '')}")
            
            # 최종 해석 추가 (강조 표시 포함)
            interpretation_lines = final.get('interpretation', '').split('\n')
            highlight_lines = final.get('highlight_lines', [])
            
            for line in interpretation_lines:
                if line.strip():
                    paragraph = doc.add_paragraph()
                    
                    # 강조할 효인지 확인
                    should_highlight = False
                    for highlight_line in highlight_lines:
                        effect_num = highlight_line + 1
                        if line.strip().startswith(f"{effect_num}효"):
                            should_highlight = True
                            break
                    
                    if should_highlight:
                        run = paragraph.add_run(line)
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
                        run.font.size = Pt(14)
                    else:
                        paragraph.add_run(line)
        
        # 파일 저장
        doc.save(filepath)
        
        # 이메일 전송
        try:
            # Gmail SMTP 설정
            smtp_server = "smtp.gmail.com"
            smtp_port = 587
            
            # 이메일 메시지 생성
            msg = MIMEMultipart()
            msg['From'] = SENDER_EMAIL
            msg['To'] = RECIPIENT_EMAIL
            msg['Subject'] = f"주역 점괘 결과 - {now.strftime('%Y년 %m월 %d일')}"
            
            # 이메일 본문
            body = f"""안녕하세요,

주역 점괘 결과를 첨부파일로 보내드립니다.

점괘 일시: {now.strftime('%Y년 %m월 %d일 %H시 %M분')}
질문: {question_text}

첨부된 워드 파일을 확인해주세요.

감사합니다.
"""
            msg.attach(MIMEText(body, 'plain', 'utf-8'))
            
            # 파일 첨부
            with open(filepath, "rb") as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
            
            encoders.encode_base64(part)
            part.add_header(
                'Content-Disposition',
                f'attachment; filename="{filename}"',
            )
            msg.attach(part)
            
            # 실제 이메일 전송
            try:
                server = smtplib.SMTP(smtp_server, smtp_port)
                server.starttls()
                server.login(SENDER_EMAIL, GMAIL_APP_PASSWORD)
                text = msg.as_string()
                server.sendmail(SENDER_EMAIL, RECIPIENT_EMAIL, text)
                server.quit()
                
                print(f"이메일 전송 성공: {filename} -> {RECIPIENT_EMAIL}")
                
                return jsonify({
                    'success': True,
                    'message': f'이메일이 {RECIPIENT_EMAIL}로 성공적으로 전송되었습니다.',
                    'filename': filename
                })
                
            except smtplib.SMTPAuthenticationError:
                return jsonify({
                    'success': False,
                    'error': 'Gmail 인증에 실패했습니다. 앱 비밀번호를 확인해주세요.',
                    'message': f'파일은 저장되었습니다: {filename}'
                })
            except smtplib.SMTPException as smtp_error:
                return jsonify({
                    'success': False,
                    'error': f'SMTP 오류: {str(smtp_error)}',
                    'message': f'파일은 저장되었습니다: {filename}'
                })
            
        except Exception as e:
            # 이메일 전송 실패 시에도 파일은 저장되었음을 알림
            return jsonify({
                'success': False,
                'error': f'이메일 전송 중 오류가 발생했습니다: {str(e)}',
                'message': f'파일은 저장되었습니다: {filename}'
            })
        
    except Exception as e:
        return jsonify({'error': f'처리 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/api/save-history', methods=['POST'])
def save_history():
    """점괘 결과를 데이터베이스에 저장 - Vercel 환경에서는 비활성화"""
    try:
        # Vercel 환경에서는 파일 시스템이 읽기 전용이므로 SQLite 사용 불가
        if os.environ.get('VERCEL'):
            return jsonify({
                'success': True,
                'message': '서버리스 환경에서는 히스토리 저장이 지원되지 않습니다.'
            })
        
        data = request.json
        
        # 데이터베이스 연결
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        # 테이블 생성
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                question TEXT,
                lines TEXT,
                moving_lines TEXT,
                original_number INTEGER,
                original_name TEXT,
                original_interpretation TEXT,
                changed_number INTEGER,
                changed_name TEXT,
                changed_interpretation TEXT,
                final_number INTEGER,
                final_name TEXT,
                final_interpretation TEXT,
                rule TEXT,
                highlight_lines TEXT,
                created_at TEXT
            )
        ''')
        
        # 데이터 저장
        question = data.get('question', '')
        lines = json.dumps(data.get('lines', []))
        moving_lines = json.dumps(data.get('moving_lines', []))
        original = data.get('original', {})
        changed = data.get('changed', {})
        final = data.get('final', {})
        rule = final.get('rule', '') if final else ''
        highlight_lines = json.dumps(final.get('highlight_lines', []) if final else [])
        created_at = datetime.now().isoformat()
        
        cursor.execute('''
            INSERT INTO history (question, lines, moving_lines, original_number, original_name, original_interpretation, changed_number, changed_name, changed_interpretation, final_number, final_name, final_interpretation, rule, highlight_lines, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (question, lines, moving_lines, original.get('number'), original.get('name'), original.get('interpretation', ''), changed.get('number'), changed.get('name'), changed.get('interpretation', ''), final.get('number'), final.get('name'), final.get('interpretation', ''), rule, highlight_lines, created_at))
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '점괘 결과가 데이터베이스에 성공적으로 저장되었습니다.'
        })
        
    except Exception as e:
        return jsonify({'error': f'히스토리 저장 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/api/history', methods=['GET'])
def get_history():
    """저장된 점괘 히스토리 조회 - Vercel 환경에서는 비활성화"""
    try:
        # Vercel 환경에서는 파일 시스템이 읽기 전용이므로 SQLite 사용 불가
        if os.environ.get('VERCEL'):
            return jsonify({
                'success': True,
                'history': [],
                'message': '서버리스 환경에서는 히스토리 기능이 지원되지 않습니다.'
            })
        
        # 데이터베이스 연결
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        # 최근 50개 결과 조회
        cursor.execute('''
            SELECT * FROM history 
            ORDER BY created_at DESC 
            LIMIT 50
        ''')
        
        rows = cursor.fetchall()
        conn.close()
        
        # 결과 포맷팅
        history = []
        for row in rows:
            history.append({
                'id': row[0],
                'question': row[1],
                'lines': json.loads(row[2]),
                'moving_lines': json.loads(row[3]),
                'original': {
                    'number': row[4],
                    'name': row[5],
                    'interpretation': row[6]
                },
                'changed': {
                    'number': row[7],
                    'name': row[8],
                    'interpretation': row[9]
                } if row[7] else None,
                'final': {
                    'number': row[10],
                    'name': row[11],
                    'interpretation': row[12]
                },
                'rule': row[13],
                'highlight_lines': json.loads(row[14]),
                'created_at': row[15]
            })
        
        return jsonify({
            'success': True,
            'history': history
        })
        
    except Exception as e:
        return jsonify({'error': f'히스토리 조회 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/api/history/<int:history_id>', methods=['DELETE'])
def delete_history_item(history_id):
    """특정 히스토리 항목 삭제 - Vercel 환경에서는 비활성화"""
    try:
        # Vercel 환경에서는 파일 시스템이 읽기 전용이므로 SQLite 사용 불가
        if os.environ.get('VERCEL'):
            return jsonify({
                'success': True,
                'message': '서버리스 환경에서는 히스토리 기능이 지원되지 않습니다.'
            })
        
        # 데이터베이스 연결
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        # 해당 ID의 히스토리 삭제
        cursor.execute('DELETE FROM history WHERE id = ?', (history_id,))
        
        if cursor.rowcount == 0:
            conn.close()
            return jsonify({'error': '해당 히스토리를 찾을 수 없습니다.'}), 404
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': '히스토리가 성공적으로 삭제되었습니다.'
        })
        
    except Exception as e:
        return jsonify({'error': f'히스토리 삭제 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/api/history/clear', methods=['DELETE'])
def clear_all_history():
    """모든 히스토리 삭제 - Vercel 환경에서는 비활성화"""
    try:
        # Vercel 환경에서는 파일 시스템이 읽기 전용이므로 SQLite 사용 불가
        if os.environ.get('VERCEL'):
            return jsonify({
                'success': True,
                'message': '서버리스 환경에서는 히스토리 기능이 지원되지 않습니다.'
            })
        
        # 데이터베이스 연결
        conn = sqlite3.connect(get_db_path())
        cursor = conn.cursor()
        
        # 모든 히스토리 삭제
        cursor.execute('DELETE FROM history')
        deleted_count = cursor.rowcount
        
        conn.commit()
        conn.close()
        
        return jsonify({
            'success': True,
            'message': f'{deleted_count}개의 히스토리가 모두 삭제되었습니다.'
        })
        
    except Exception as e:
        return jsonify({'error': f'히스토리 전체 삭제 중 오류가 발생했습니다: {str(e)}'}), 500

@app.route('/history')
def history_page():
    """히스토리 페이지"""
    return render_template('history.html')

if __name__ == '__main__':
    # 필요한 폴더 생성 (로컬 환경에서만)
    if not os.environ.get('VERCEL'):
        Path("해석").mkdir(exist_ok=True)
        Path("출력").mkdir(exist_ok=True)
    
    app.run(host='0.0.0.0', port=5002, debug=True)

# Vercel용 WSGI 애플리케이션
application = app 